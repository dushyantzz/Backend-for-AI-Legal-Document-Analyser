from models import Document, Template, db, DocumentType
from datetime import datetime
import os
import mammoth
from docx import Document as DocxDocument
import requests

class DocumentService:
    def __init__(self):
        pass
    
    def create_document(self, user_id, document_data):
        """Create a new document"""
        try:
            document = Document(
                user_id=user_id,
                title=document_data['title'],
                document_type=DocumentType(document_data['document_type']),
                template_id=document_data.get('template_id'),
                content=document_data.get('content'),
                metadata=document_data.get('metadata', {})
            )
            
            db.session.add(document)
            db.session.commit()
            
            return document
            
        except Exception as e:
            db.session.rollback()
            raise e
    
    def get_user_documents(self, user_id, page=1, per_page=10, document_type=None):
        """Get documents for a user with pagination"""
        try:
            query = Document.query.filter_by(user_id=user_id, is_active=True)
            
            if document_type:
                query = query.filter_by(document_type=DocumentType(document_type))
            
            documents = query.paginate(
                page=page, per_page=per_page, error_out=False
            )
            
            return documents
            
        except Exception as e:
            raise e
    
    def get_document(self, document_id, user_id):
        """Get a specific document"""
        try:
            document = Document.query.filter_by(
                id=document_id, user_id=user_id, is_active=True
            ).first()
            
            if not document:
                raise ValueError('Document not found')
            
            return document
            
        except Exception as e:
            raise e
    
    def update_document(self, document_id, user_id, update_data):
        """Update document content"""
        try:
            document = Document.query.filter_by(
                id=document_id, user_id=user_id, is_active=True
            ).first()
            
            if not document:
                raise ValueError('Document not found')
            
            # Update allowed fields
            if 'title' in update_data:
                document.title = update_data['title']
            if 'content' in update_data:
                document.content = update_data['content']
            if 'metadata' in update_data:
                document.metadata = update_data['metadata']
            
            document.version += 1
            document.updated_at = datetime.utcnow()
            db.session.commit()
            
            return document
            
        except Exception as e:
            db.session.rollback()
            raise e
    
    def delete_document(self, document_id, user_id):
        """Soft delete a document"""
        try:
            document = Document.query.filter_by(
                id=document_id, user_id=user_id, is_active=True
            ).first()
            
            if not document:
                raise ValueError('Document not found')
            
            document.is_active = False
            document.updated_at = datetime.utcnow()
            db.session.commit()
            
            return True
            
        except Exception as e:
            db.session.rollback()
            raise e
    
    def generate_document_from_template(self, template_id, form_data):
        """Generate document from template with form data"""
        try:
            template = Template.query.get(template_id)
            
            if not template:
                raise ValueError('Template not found')
            
            # Replace placeholders in template content
            content = template.template_content
            
            # Replace form data placeholders
            for key, value in form_data.items():
                placeholder = f"#{key}#"
                content = content.replace(placeholder, str(value))
            
            return content
            
        except Exception as e:
            raise e
    
    def download_document(self, document_id, user_id):
        """Download document as file"""
        try:
            document = Document.query.filter_by(
                id=document_id, user_id=user_id, is_active=True
            ).first()
            
            if not document:
                raise ValueError('Document not found')
            
            # If document has file_path, return that file
            if document.file_path and os.path.exists(document.file_path):
                return document.file_path
            
            # Otherwise, create a temporary file from content
            temp_path = f"/tmp/document_{document_id}.docx"
            
            # Create a new document
            doc = DocxDocument()
            
            # Add content to document
            if document.content:
                doc.add_paragraph(document.content)
            
            # Save document
            doc.save(temp_path)
            
            return temp_path
            
        except Exception as e:
            raise e
    
    def convert_docx_to_html(self, docx_path):
        """Convert DOCX file to HTML"""
        try:
            with open(docx_path, 'rb') as docx_file:
                result = mammoth.convert_to_html(docx_file)
                return result.value
                
        except Exception as e:
            raise e
    
    def get_templates(self, document_type=None):
        """Get available templates"""
        try:
            query = Template.query.filter_by(is_active=True)
            
            if document_type:
                query = query.filter_by(document_type=DocumentType(document_type))
            
            templates = query.all()
            return templates
            
        except Exception as e:
            raise e
    
    def create_template(self, template_data, created_by):
        """Create a new template (admin only)"""
        try:
            template = Template(
                name=template_data['name'],
                document_type=DocumentType(template_data['document_type']),
                description=template_data.get('description'),
                template_content=template_data['template_content'],
                form_schema=template_data['form_schema'],
                created_by=created_by
            )
            
            db.session.add(template)
            db.session.commit()
            
            return template
            
        except Exception as e:
            db.session.rollback()
            raise e
